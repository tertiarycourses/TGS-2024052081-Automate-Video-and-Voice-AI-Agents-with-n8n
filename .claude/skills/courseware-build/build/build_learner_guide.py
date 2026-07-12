#!/usr/bin/env python3
"""Generate the CompTIA PenTest+ (PT0-003) Learner Guide as BOTH a Markdown mirror (LG-*.md at repo
root) and a DOCX (courseware/LG-*.docx) from one source, so they never diverge.

House format: cover page, Document Version Control Record, auto TOC, Arial 11pt
body, one section per lab (Objective · Goal · What you'll build · Step-by-step
with commands · Test it), plus setup, exam-prep and glossary. All content is
driven by course_data + the domain data files, keeping the LG 100% aligned with
the slide deck, Lesson Plan and labs.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3; from data_domain4 import DOMAIN4
from data_domain5 import DOMAIN5
ACT=DOMAIN1+DOMAIN2+DOMAIN3+DOMAIN4+DOMAIN5
import prodoc
def _find_repo(start):
    env=os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d=start
    for _ in range(8):
        d=os.path.dirname(d)
        if os.path.isdir(os.path.join(d,"courseware")) and os.path.isdir(os.path.join(d,"labs")): return d
    return os.path.dirname(os.path.dirname(HERE))
REPO=_find_repo(HERE); ASSETS=os.path.join(os.path.dirname(HERE),"assets")

# ---------------- block DSL (single content stream → MD + DOCX) ----------------
B=[]
def h1(t): B.append(("h1",t))
def h2(t): B.append(("h2",t))
def h3(t): B.append(("h3",t))
def p(t):  B.append(("p",t))
def bullets(xs): B.append(("bullets",xs))
def steps(xs): B.append(("steps",xs))
def code(t): B.append(("code",t))
def note(t): B.append(("note",t))
def rule(): B.append(("rule",))

# ---------------- content ----------------
h1("Introduction")
p(f"This Learner Guide accompanies the WSQ course {C.TITLE} ({C.COURSE_CODE}), conducted by {C.ORG}. "
  "It provides step-by-step instructions for all 37 hands-on labs, organised by the five official "
  "PT0-003 exam domains. Every lab maps to a published exam objective and is completed on Kali Linux "
  "using industry-standard, free penetration-testing tools.")
p("Use this guide alongside the course slides and the lab files in the labs/ folder of the course "
  "repository. All offensive techniques must only be run against systems you own or have written "
  "authorisation to test — unauthorised access to a computer is a criminal offence under the "
  "Singapore Computer Misuse Act and equivalents elsewhere. Always complete the clean-up and "
  "artifact-preservation steps at the end of an engagement.")

h1("Course Learning Outcomes")
bullets(C.LEARNING_OUTCOMES)

h1("Before You Start — Environment Setup")
h3("What you need")
bullets([
 "A Kali Linux environment — the Killercoda Kali playground (browser-based), a Kali VM in VirtualBox/VMware, Kali in WSL2, or the Kali Docker image.",
 "Most tools are pre-installed in Kali; install anything missing with apt (e.g. sudo apt install -y <tool>).",
 "Vulnerable practice targets launched as Docker containers on the Kali host: DVWA, OWASP Juice Shop and Metasploitable3; plus AWS Free Tier / localstack for the cloud lab.",
 "Written authorisation (a signed Rules of Engagement) for any target that is not your own lab machine.",
])
h3("Launch and verify Kali Linux")
p("Kali Linux ships with the offensive-security toolchain pre-installed (Nmap, Metasploit, Burp Suite, hashcat, Aircrack-ng and hundreds more). Start your Kali environment, open a terminal, become root when a step needs it, and confirm the core tools are present before you begin.")
code("kali@kali:~$ id                 # confirm your user\nkali@kali:~$ nmap --version      # core scanner present\nkali@kali:~$ msfconsole --version # Metasploit present\nkali@kali:~$ sudo apt update      # refresh package lists")
h3("Conventions used in every lab")
bullets([
 "Commands are run from a Kali terminal; a leading sudo is used where root is required.",
 "Placeholders such as <TARGET>, <KALI> and <ip> are replaced with your own lab addresses.",
 "Each lab creates a working directory under ~/engagements/ for its notes, output and loot.",
 "Reset or re-deploy a vulnerable target between exploit labs that change its state.",
 "Run the clean-up step at the end of state-changing labs, and preserve artifacts for the report.",
])

# ---------------- per-topic, per-lab ----------------
TOPICS_BY_NUM={t["num"]:t for t in C.TOPICS}
for t in C.TOPICS:
    h1(f"Topic {t['code']} — {t['title']}  ({t['weighting']})")
    p(t["subtitle"])
    h3("Key concepts")
    bullets(t["concepts"])
    for a in [x for x in ACT if x["topic"]==t["num"]]:
        h2(f"Lab {a['num']} — {a['title']}")
        p(f"Exam objective: {a['objective']}.")
        p(f"Goal: {a['desc']}")
        h3("What you'll build")
        p(a["build"]+f"   (Tools: {a['services']}.)")
        h3("Step-by-step")
        st=[]
        for i,(instr,cmd) in enumerate(a["steps"],1):
            st.append((instr,cmd))
        steps(st)
        h3("Test it")
        p(a["test"])
        note(f"Full commands and screenshots are in labs/lab-{a['num']:02d}-*.md. "
             f"Only run these techniques against systems you own or are authorised to test.")
        rule()

h1("Exam Focus — Testing Frameworks, Scan Types and Specialized Systems")
p("These cross-cutting topics are examined throughout PT0-003 but do not each have a dedicated hands-on lab. "
  "Study this section alongside the labs so you can answer the knowledge questions and the specialized-systems items on the exam.")
h3("Testing frameworks and methodologies (Objective 1.3)")
p("A tester chooses a recognised methodology to make the engagement systematic, repeatable and defensible. Know what each is for:")
bullets([
 "PTES (Penetration Testing Execution Standard) — the seven-phase pentest process (pre-engagement → intelligence gathering → threat modelling → vulnerability analysis → exploitation → post-exploitation → reporting).",
 "OSSTMM — the Open Source Security Testing Methodology Manual, for measurable, repeatable testing.",
 "OWASP Top 10 + WSTG (Web Security Testing Guide) for web apps, and OWASP MASVS for mobile apps.",
 "MITRE ATT&CK — a knowledge base of adversary tactics and techniques used to give findings a threat-realistic narrative.",
 "NIST SP 800-115 and CREST — recognised technical testing and accreditation standards.",
 "STRIDE and DREAD — threat-modelling frameworks; the Purdue model structures OT/ICS networks into levels.",
])
h3("Scan types — SAST, DAST, IAST and SCA (Objective 3.1)")
bullets([
 "SAST (Static Application Security Testing) — analyses source code WITHOUT running it; used early in the SDLC, finds code-level flaws but yields more false positives.",
 "DAST (Dynamic Application Security Testing) — tests the RUNNING application from the outside (e.g. OWASP ZAP, Nikto); finds runtime and configuration issues but has no code visibility.",
 "IAST (Interactive AST) — instruments the running application to combine static and dynamic visibility.",
 "SCA (Software Composition Analysis) — inventories third-party/open-source dependencies and flags known-CVE components (e.g. Trivy, Grype).",
 "Scans may be authenticated or unauthenticated, and extend to Infrastructure-as-Code and container images for full-stack coverage.",
])
h3("Specialized and emerging systems (Objective 4.9)")
bullets([
 "Mobile — analyse apps with MobSF (static/dynamic), Frida (runtime instrumentation), Drozer and ADB; typical weaknesses are insecure data storage, hardcoded secrets/API keys, weak or missing certificate pinning and over-broad permissions; jailbreak/root to bypass client-side controls.",
 "IoT / OT / ICS — fragile protocols with little or no authentication (Modbus, DNP3, CAN bus); test with traffic capture, replay and register manipulation, and map the network with the Purdue model — but NEVER disrupt a safety-critical process; prefer passive/read-only techniques.",
 "AI / ML systems (new in PT0-003) — prompt injection (direct and indirect), model manipulation and data poisoning, jailbreaks / guardrail bypass, sensitive- or training-data extraction, and an insecure model supply chain.",
 "RF / NFC — wardriving, RFID/badge cloning, Bluejacking and Bluetooth attacks against physical and wireless entry points.",
 "Across all specialized systems, stay strictly within scope and the Rules of Engagement — they are fragile and often safety-relevant.",
])
rule()

h1("Exam Preparation")
bullets([
 "First pass: complete every lab on Kali Linux, reading the References in each lab file.",
 "Second pass: redo the labs from memory until the tool workflow and flags are automatic.",
 "Review the 'Test it' check and the 'What you learned' bullets for any topic you find hard.",
 "Practise mapping each finding to a CVSS score, a CWE weakness and a MITRE ATT&CK technique.",
 "Sharpen exam readiness with the Tertiary Infotech PenTest+ practice exam: https://exams.tertiaryinfotech.com/practice-exams/comptia/comptia-pentest-plus",
 "Take the free CompTIA practice assessment for PT0-003 and sit the exam via a Pearson VUE test centre or online proctoring.",
])

h1("Glossary")
gl=[
 ("Rules of Engagement (RoE)","The signed authorisation defining scope, windows, allowed/forbidden actions and stop conditions."),
 ("Statement of Work (SoW)","The contract defining targets, methodology, schedule, deliverables and fees for the engagement."),
 ("OSINT","Open-Source Intelligence — information gathered from public sources without touching the target."),
 ("Enumeration","Actively probing a target to list hosts, ports, services, users and shares."),
 ("CVSS","Common Vulnerability Scoring System — a 0–10 severity score for a vulnerability."),
 ("CWE / EPSS","Common Weakness Enumeration (weakness type) / Exploit Prediction Scoring System (exploit likelihood)."),
 ("MITRE ATT&CK","A knowledge base of adversary tactics and techniques used to map findings to real attacks."),
 ("Payload","The code a delivered exploit runs — e.g. a reverse shell or Meterpreter session."),
 ("Privilege escalation","Turning a low-privileged foothold into administrator/root control of a host."),
 ("Lateral movement","Using captured credentials or hashes to move from one host to another across a network."),
 ("Pivoting","Routing traffic through a compromised host to reach otherwise-unreachable internal networks."),
 ("Persistence","Mechanisms (cron, services, keys, web shells) that keep access across reboots."),
]
B.append(("dl",gl))

# ---------------- render Markdown ----------------
def _anchor(txt):
    return "".join(ch.lower() if ch.isalnum() else ("-" if ch in " -" else "") for ch in txt)

def render_md():
    out=[f"# {C.TITLE} — Learner Guide",""]
    out.append(f"**WSQ Course Code:** {C.COURSE_CODE}  |  **Conducted by:** {C.ORG} ({C.UEN.replace('UEN: ','UEN ')})  |  **Version {C.VERSION} · {C.VERSION_DATE}**")
    out.append("")
    # TOC (h1 + h2)
    out.append("## Contents"); out.append("")
    for kind,*rest in B:
        if kind=="h1": out.append(f"- [{rest[0]}](#{_anchor(rest[0])})")
        elif kind=="h2": out.append(f"  - [{rest[0]}](#{_anchor(rest[0])})")
    out.append("")
    for kind,*rest in B:
        if kind=="h1": out+=["",f"## {rest[0]}",""]
        elif kind=="h2": out+=["",f"### {rest[0]}",""]
        elif kind=="h3": out+=[f"**{rest[0]}**",""]
        elif kind=="p": out+=[rest[0],""]
        elif kind=="bullets": out+=[f"- {x}" for x in rest[0]]+[""]
        elif kind=="steps":
            for i,(instr,cmd) in enumerate(rest[0],1):
                out.append(f"{i}. {instr}")
                if cmd: out+=["",f"   ```bash",f"   {cmd}","   ```",""]
            out.append("")
        elif kind=="code": out+=["```bash",rest[0],"```",""]
        elif kind=="note": out+=[f"> **Note:** {rest[0]}",""]
        elif kind=="rule": out+=["---",""]
        elif kind=="dl":
            for term,defn in rest[0]: out.append(f"- **{term}** — {defn}")
            out.append("")
    return "\n".join(out)

MD_OUT=os.path.join(REPO,f"LG-{C.SHORT_TITLE}.md")
with open(MD_OUT,"w") as f: f.write(render_md())
print("Saved",MD_OUT)

# ---------------- render DOCX ----------------
BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
INKCODE=RGBColor(0x0B,0x30,0x60)
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)
prodoc.add_cover_page(doc,"LEARNER GUIDE",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[
 ("1.0",C.VERSION_DATE,"Initial release — CompTIA PenTest+ (PT0-003) Learner Guide covering all 35 labs.",C.TRAINER),
 (C.VERSION.lstrip("v"),C.VERSION_DATE,"Added Lab 12 (Wireshark/tcpdump traffic analysis) and Lab 21 (Netcat/Socat shells) — now 37 labs; added an Exam Focus section on testing frameworks, scan types (SAST/DAST/IAST/SCA) and specialized systems (mobile/IoT/OT/AI) for full PT0-003 domain alignment.",C.TRAINER),
])
prodoc.add_toc(doc)

def code_para(text):
    for line in text.split("\n"):
        para=doc.add_paragraph(); prodoc._shade_para(para) if hasattr(prodoc,"_shade_para") else None
        r=para.add_run(line); r.font.name="Consolas"; r.font.size=Pt(9.5); r.font.color.rgb=INKCODE

for kind,*rest in B:
    if kind=="h1": doc.add_heading(rest[0],level=1)
    elif kind=="h2": doc.add_heading(rest[0],level=2)
    elif kind=="h3":
        para=doc.add_paragraph(); r=para.add_run(rest[0]); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=BRAND
    elif kind=="p": doc.add_paragraph(rest[0])
    elif kind=="bullets":
        for x in rest[0]: doc.add_paragraph(x,style="List Bullet")
    elif kind=="steps":
        for i,(instr,cmd) in enumerate(rest[0],1):
            para=doc.add_paragraph(style="List Number"); para.add_run(instr)
            if cmd: code_para(cmd)
    elif kind=="code": code_para(rest[0])
    elif kind=="note":
        para=doc.add_paragraph(); r=para.add_run("Note: "); r.bold=True; r.font.color.rgb=BRAND
        para.add_run(rest[0]).font.size=Pt(10)
    elif kind=="rule": doc.add_paragraph("")
    elif kind=="dl":
        for term,defn in rest[0]:
            para=doc.add_paragraph(style="List Bullet")
            r=para.add_run(term+" — "); r.bold=True; para.add_run(defn)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
DOCX_OUT=os.path.join(REPO,"courseware",f"LG-{C.SHORT_TITLE}.docx")
doc.save(DOCX_OUT)
print("Saved",DOCX_OUT)
